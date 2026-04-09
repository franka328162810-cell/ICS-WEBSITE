"""Extract Chapter 27 (Comprehensive Stress Test Report) key content from Volume 3."""
import docx, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = docx.Document(r'C:\Users\Administrator\Desktop\《星际文明的规范性基础：新三观与深不确定性下的治理》\9.0书籍提交\2025-1-25第三卷定稿.docx')
paras = doc.paragraphs

# Chapter 27 spans roughly para 13665 to 15948 (before Ch28)
# Key tables identified from markers:
# 表27.3 AGI情境测试主要发现 (para ~13779)
# 表27.5 AGI情境修订建议汇总 (para ~13809)
# 表27.7 FC情境测试主要发现 (para ~13848)
# 表27.9 FC情境修订建议汇总 (para ~13877)
# 表27.11 生态临界点情境测试主要发现 (para ~13916)
# 表27.13 生态临界点情境修订建议汇总 (para ~13945)
# 表27.15 代际正义危机情境测试主要发现 (para ~13984)
# 表27.17 代际正义危机情境修订建议汇总 (para ~14014)
# 表27.x 鲁棒性得分 (para ~14709)
# 表27.50 压力测试元反思结论汇总 (para ~15195)

# Also extract tables from doc.tables that fall within Ch27 content
# First, let's extract the text content of Ch27 key sections

out = []

def add(s):
    out.append(s)

# Extract paragraphs 13665-15315 (main Ch27 content before footnotes)
add("=" * 80)
add("第二十七章 压力测试综合报告与理论修订")
add("Chapter 27: Comprehensive Stress Test Report and Theory Revision")
add("=" * 80)

# Scan for section headers and key content in Ch27 range
ch27_start = 13665
ch27_end = 15315  # Before footnotes

for i in range(ch27_start, min(ch27_end, len(paras))):
    p = paras[i]
    text = p.text.strip()
    if not text:
        continue
    
    # Check if bold (section header)
    is_bold = any(run.bold for run in p.runs if run.text.strip())
    
    # Include all non-empty paragraphs in key ranges
    # But especially mark tables and headers
    if is_bold:
        add(f"\n[BOLD] [{i}] {text}")
    else:
        add(f"[{i}] {text}")

# Now extract the actual Word tables in Ch27 region
# We need to find tables by checking their position relative to paragraphs
add("\n\n" + "=" * 80)
add("WORD TABLES IN CHAPTER 27 (extracted from doc.tables)")
add("=" * 80)

# Get all tables and try to identify Ch27 ones by content
for t_idx, table in enumerate(doc.tables):
    # Check first cell content
    try:
        first_cell = table.rows[0].cells[0].text.strip()
    except:
        continue
    
    # Look for Ch27 specific tables
    ch27_markers = ['框架组件', '评估层面', '弱点类型', '修订类型', '鲁棒性', 
                    '反思维度', '分析维度', '情境', 'AGI', 'First Contact',
                    '生态临界点', '代际正义', '综合评分']
    
    # Check if any row mentions "27" or key Ch27 terms
    all_text = ''
    for row in table.rows[:3]:
        for cell in row.cells:
            all_text += cell.text + ' '
    
    is_ch27_table = False
    if '27.' in all_text or '表27' in all_text:
        is_ch27_table = True
    elif any(m in all_text for m in ['鲁棒性得分', '鲁棒性等级', '修订建议汇总', '弱点汇总']):
        is_ch27_table = True
    elif any(m in first_cell for m in ch27_markers):
        # Additional check - see if it's in the right position
        # Check if table content matches Ch27 patterns
        for row in table.rows:
            for cell in row.cells:
                if '27' in cell.text:
                    is_ch27_table = True
                    break
            if is_ch27_table:
                break
    
    if is_ch27_table:
        add(f"\n--- Table #{t_idx} ---")
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            add(f"  Row {r_idx}: {' | '.join(cells)}")
        if len(table.rows) > 30:
            add(f"  ... (total {len(table.rows)} rows)")

content = '\n'.join(out)
with open(r'c:\Users\Administrator\Documents\ics-website\vol3_ch27_raw.txt', 'w', encoding='utf-8') as f:
    f.write(content)

print(f"Extracted {len(out)} lines to vol3_ch27_raw.txt")
print(f"File size: {len(content)} chars")
